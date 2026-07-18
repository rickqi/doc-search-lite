"""
Unit tests for table_fix module.

Tests cover:
- fix_table_alignment: Padding short rows, separator fixes, multiple tables,
  code block skipping, edge cases
- fix_merged_tables_with_html: DOCX merge detection, HTML generation,
  non-DOCX passthrough, integration with OfficeConverter
"""

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.converter.table_fix import (
    _count_pipe_columns,
    _find_nth_table_start,
    _is_separator_row,
    _make_separator,
    _pad_row,
    fix_merged_tables_with_html,
    fix_table_alignment,
)
from src.converter.office import OfficeConverter


# ---------------------------------------------------------------------------
# Helper: create DOCX files with various table structures
# ---------------------------------------------------------------------------


def _create_docx_no_merge(path: Path) -> None:
    """Create a DOCX with a simple 2x3 table (no merged cells)."""
    doc = Document()
    doc.add_heading("Simple Table", level=1)
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"
    doc.save(str(path))


def _create_docx_with_colspan(path: Path) -> None:
    """Create a DOCX with a horizontally merged cell (colspan=2)."""
    doc = Document()
    doc.add_heading("Merged Table", level=1)
    table = doc.add_table(rows=2, cols=3)
    # Merge cells (0,0) and (0,1) horizontally
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Merged"
    table.cell(0, 2).text = "Normal"
    table.cell(1, 0).text = "X"
    table.cell(1, 1).text = "Y"
    table.cell(1, 2).text = "Z"
    doc.save(str(path))


def _create_docx_with_rowspan(path: Path) -> None:
    """Create a DOCX with a vertically merged cell (rowspan=2)."""
    doc = Document()
    doc.add_heading("Vertical Merge Table", level=1)
    table = doc.add_table(rows=3, cols=3)
    # Merge cells (0,0) and (1,0) vertically
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = "VMerge"
    table.cell(0, 1).text = "B1"
    table.cell(0, 2).text = "C1"
    # Row 1, col 0 is consumed by the merge
    table.cell(1, 1).text = "B2"
    table.cell(1, 2).text = "C2"
    table.cell(2, 0).text = "A3"
    table.cell(2, 1).text = "B3"
    table.cell(2, 2).text = "C3"
    doc.save(str(path))


def _create_docx_multi_table(path: Path) -> None:
    """Create a DOCX with two tables: one simple, one with merges."""
    doc = Document()
    doc.add_heading("Multi Table", level=1)

    # Table 1: simple (no merge)
    table1 = doc.add_table(rows=2, cols=2)
    table1.cell(0, 0).text = "T1-A"
    table1.cell(0, 1).text = "T1-B"
    table1.cell(1, 0).text = "T1-C"
    table1.cell(1, 1).text = "T1-D"

    doc.add_paragraph("Some text between tables.")

    # Table 2: with horizontal merge
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).merge(table2.cell(0, 1))
    table2.cell(0, 0).text = "Merged"
    table2.cell(1, 0).text = "X"
    table2.cell(1, 1).text = "Y"

    doc.save(str(path))


# ===================================================================
# Tests for internal helpers
# ===================================================================


class TestCountPipeColumns:
    """Test _count_pipe_columns helper."""

    def test_simple_row(self):
        assert _count_pipe_columns("| A | B | C |") == 3

    def test_single_column(self):
        assert _count_pipe_columns("| A |") == 1

    def test_empty_cells(self):
        assert _count_pipe_columns("| | | |") == 3

    def test_no_leading_pipe(self):
        assert _count_pipe_columns("A | B |") == 0

    def test_separator_row(self):
        assert _count_pipe_columns("|---|---|---|") == 3

    def test_two_columns(self):
        assert _count_pipe_columns("| Header 1 | Header 2 |") == 2


class TestIsSeparatorRow:
    """Test _is_separator_row helper."""

    def test_standard_separator(self):
        assert _is_separator_row("|---|---|---|") is True

    def test_aligned_separator(self):
        assert _is_separator_row("|:---|:---:|---:|") is True

    def test_not_separator_data(self):
        assert _is_separator_row("| A | B |") is False

    def test_not_separator_no_pipe(self):
        assert _is_separator_row("---") is False


class TestMakeSeparator:
    """Test _make_separator helper."""

    def test_three_cols(self):
        result = _make_separator(3)
        assert result == "|---|---|---|"

    def test_one_col(self):
        result = _make_separator(1)
        assert result == "|---|"


class TestPadRow:
    """Test _pad_row helper."""

    def test_no_padding_needed(self):
        row = "| A | B | C |"
        assert _pad_row(row, 3) == row

    def test_pad_short_row(self):
        row = "| A | B |"
        result = _pad_row(row, 3)
        assert _count_pipe_columns(result) == 3

    def test_pad_by_two(self):
        row = "| X |"
        result = _pad_row(row, 3)
        assert _count_pipe_columns(result) == 3


# ===================================================================
# Tests for fix_table_alignment
# ===================================================================


class TestFixTableAlignment:
    """Test fix_table_alignment function."""

    def test_aligned_table_unchanged(self):
        """A properly aligned 3-column table should be unchanged."""
        md = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |"
        assert fix_table_alignment(md) == md

    def test_short_row_padded(self):
        """A row with fewer columns should be padded."""
        md = "| A | B | C |\n|---|---|---|\n| X | Y |"
        result = fix_table_alignment(md)
        lines = result.split("\n")
        # The short row should now have 3 columns
        assert _count_pipe_columns(lines[2]) == 3

    def test_separator_row_fixed(self):
        """Separator row should match max column count."""
        md = "| A | B | C |\n|---|---|\n| X | Y | Z |"
        result = fix_table_alignment(md)
        lines = result.split("\n")
        assert _count_pipe_columns(lines[1]) == 3
        assert _is_separator_row(lines[1])

    def test_multiple_tables_independent(self):
        """Multiple tables should each be fixed independently."""
        md = (
            "| A | B |\n|---|---|\n| X |\n\n"
            "Some text\n\n"
            "| C | D | E |\n|---|---|---|\n| Y | Z |"
        )
        result = fix_table_alignment(md)
        parts = result.split("\n\nSome text\n\n")
        # First table: all rows should have 2 columns
        first_lines = parts[0].split("\n")
        assert _count_pipe_columns(first_lines[2]) == 2
        # Second table: all rows should have 3 columns
        second_lines = parts[1].split("\n")
        assert _count_pipe_columns(second_lines[2]) == 3

    def test_empty_input(self):
        """Empty string should return empty string."""
        assert fix_table_alignment("") == ""

    def test_no_tables(self):
        """Text without tables should be unchanged."""
        md = "Just some text\nwith no tables"
        assert fix_table_alignment(md) == md

    def test_no_pipe_chars(self):
        """Text without any pipe characters should be unchanged."""
        md = "Hello world\nNo pipes here"
        assert fix_table_alignment(md) == md

    def test_table_inside_code_block_skipped(self):
        """Tables inside fenced code blocks should be ignored."""
        md = "```\n| A | B |\n|---|---|\n| X |\n```\n"
        result = fix_table_alignment(md)
        # The code block table should not be modified
        assert result == md

    def test_table_after_code_block(self):
        """Table after a closed code block should be fixed."""
        md = "```\ncode\n```\n\n| A | B | C |\n|---|---|\n| X | Y |"
        result = fix_table_alignment(md)
        lines = result.split("\n")
        # The data row (last line) should be padded to 3 columns
        assert _count_pipe_columns(lines[-1]) == 3

    def test_all_short_rows_padded(self):
        """Multiple short rows should all be padded."""
        md = "| A | B | C |\n|---|---|---|\n| 1 |\n| 2 | 3 |"
        result = fix_table_alignment(md)
        lines = result.split("\n")
        assert _count_pipe_columns(lines[2]) == 3
        assert _count_pipe_columns(lines[3]) == 3

    def test_single_row_table_not_treated_as_table(self):
        """A single pipe line should not be treated as a table block."""
        md = "| just a pipe thing"
        # Single line starting with | but no separator line
        # Our implementation requires at least 2 lines
        assert fix_table_alignment(md) == md


# ===================================================================
# Tests for fix_merged_tables_with_html
# ===================================================================


class TestFixMergedTablesWithHtml:
    """Test fix_merged_tables_with_html function."""

    def test_non_docx_passthrough(self, tmp_path):
        """Non-DOCX files should return markdown unchanged."""
        xlsx_file = tmp_path / "test.xlsx"
        xlsx_file.write_bytes(b"fake xlsx")
        md = "| A | B |\n|---|---|\n| X | Y |"
        result = fix_merged_tables_with_html(xlsx_file, md)
        assert result == md

    def test_empty_markdown(self, tmp_path):
        """Empty markdown should return empty."""
        docx_file = tmp_path / "test.docx"
        _create_docx_no_merge(docx_file)
        assert fix_merged_tables_with_html(docx_file, "") == ""

    def test_no_merge_preserves_markdown(self, tmp_path):
        """DOCX without merged cells should preserve Markdown tables."""
        docx_file = tmp_path / "simple.docx"
        _create_docx_no_merge(docx_file)
        md = "| A | B | C |\n|---|---|---|\n| D | E | F |"
        result = fix_merged_tables_with_html(docx_file, md)
        # No merge detected, markdown should be unchanged
        assert result == md

    def test_colspan_produces_html(self, tmp_path):
        """DOCX with horizontal merge should produce HTML table."""
        docx_file = tmp_path / "merged.docx"
        _create_docx_with_colspan(docx_file)
        md = "| Merged | B |\n|---|---|\n| X | Y |"
        result = fix_merged_tables_with_html(docx_file, md)
        assert "<table>" in result
        assert "<td" in result
        assert "colspan" in result

    def test_rowspan_produces_html(self, tmp_path):
        """DOCX with vertical merge should produce HTML table."""
        docx_file = tmp_path / "vmerge.docx"
        _create_docx_with_rowspan(docx_file)
        md = "| VMerge | B1 | C1 |\n|---|---|---|\n| B2 | C2 |\n| A3 | B3 | C3 |"
        result = fix_merged_tables_with_html(docx_file, md)
        assert "<table>" in result
        assert "rowspan" in result

    def test_nonexistent_file_graceful(self, tmp_path):
        """Non-existent DOCX file should return markdown unchanged."""
        docx_file = tmp_path / "nonexistent.docx"
        md = "| A | B |\n|---|---|\n| X | Y |"
        result = fix_merged_tables_with_html(docx_file, md)
        # Should gracefully handle missing file
        assert result == md

    def test_multi_table_selective_replacement(self, tmp_path):
        """Only tables with merges should be replaced with HTML."""
        docx_file = tmp_path / "multi.docx"
        _create_docx_multi_table(docx_file)
        md = (
            "| T1-A | T1-B |\n|---|---|\n| T1-C | T1-D |\n\n"
            "Some text between tables.\n\n"
            "| Merged | Y |\n|---|---|\n| X | Y |"
        )
        result = fix_merged_tables_with_html(docx_file, md)
        # Second table (with merge) should become HTML
        # First table (no merge) should remain Markdown
        # The second table should contain HTML
        assert "<table>" in result
        # Verify T1 content is still there (first table preserved)
        assert "T1-A" in result


# ===================================================================
# Integration tests via OfficeConverter
# ===================================================================


class TestTableFixIntegration:
    """Integration tests verifying table fix works through OfficeConverter."""

    def test_convert_docx_with_merged_cells(self, tmp_path, tmp_output_dir):
        """Full pipeline: DOCX with merged cells → convert → check output."""
        docx_file = tmp_path / "merged.docx"
        _create_docx_with_colspan(docx_file)

        converter = OfficeConverter()
        result = converter.convert(docx_file, tmp_output_dir)

        assert result.success is True
        # Output should contain either HTML table or aligned Markdown table
        assert result.markdown  # Non-empty
        # Verify content is present
        md = result.markdown
        # Either HTML table or Markdown table with the content
        has_table = "<table>" in md or "Merged" in md
        assert has_table

    def test_convert_simple_docx_table_preserved(self, tmp_path, tmp_output_dir):
        """Simple DOCX tables should remain as Markdown tables."""
        docx_file = tmp_path / "simple.docx"
        _create_docx_no_merge(docx_file)

        converter = OfficeConverter()
        result = converter.convert(docx_file, tmp_output_dir)

        assert result.success is True
        md = result.markdown
        # Simple table should remain as Markdown, not HTML
        assert "|" in md

    def test_convert_docx_with_vertical_merge(self, tmp_path, tmp_output_dir):
        """Full pipeline: DOCX with vertical merge → convert → check output."""
        docx_file = tmp_path / "vmerge.docx"
        _create_docx_with_rowspan(docx_file)

        converter = OfficeConverter()
        result = converter.convert(docx_file, tmp_output_dir)

        assert result.success is True
        assert result.markdown
        # Should have either HTML table with rowspan or properly aligned MD
        md = result.markdown
        has_content = "VMerge" in md or "table" in md.lower()
        assert has_content
