"""
Table alignment fix and HTML fallback for merged cells.

When converting DOCX files with merged cells (rowspan/colspan) to Markdown,
tables become misaligned because markdownify ignores rowspan attributes.
This module provides two fixes:

1. fix_table_alignment: Pads short rows in Markdown tables to consistent column counts
2. fix_merged_tables_with_html: Replaces merged-cell Markdown tables with HTML <table>
"""

import logging
import re
from pathlib import Path
from typing import List, Optional, Tuple
from xml.etree import ElementTree as ET

logger = logging.getLogger(__name__)

# Namespace for DOCX XML
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _count_pipe_columns(line: str) -> int:
    """Count the number of columns in a Markdown table row.

    A row like ``| A | B | C |`` has 3 columns.
    Handles edge cases like empty cells and trailing pipes.
    """
    stripped = line.strip()
    if not stripped.startswith("|"):
        return 0
    # Split by | and count segments between pipes
    parts = stripped.split("|")
    # First and last elements are empty strings from leading/trailing |
    # e.g. "| A | B |" -> ['', ' A ', ' B ', '']
    # Filter to only the data segments
    if stripped.endswith("|"):
        # Remove first and last empty strings
        data_parts = parts[1:-1]
    else:
        # Trailing | missing: "| A | B " -> ['', ' A ', ' B ']
        data_parts = parts[1:]
    return len(data_parts)


def _is_separator_row(line: str) -> bool:
    """Check if a line is a Markdown table separator row (e.g. |---|---|)."""
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    # Remove leading and trailing pipes
    inner = stripped.strip("|")
    # Each segment should be dashes/colons (e.g. ---, :---, ---:, :---:)
    segments = inner.split("|")
    return all(re.match(r"^:?-+:?$", s.strip()) for s in segments)


def _make_separator(num_cols: int) -> str:
    """Create a separator row for the given number of columns."""
    return "|" + "|".join(["---"] * num_cols) + "|"


def _pad_row(line: str, target_cols: int) -> str:
    """Pad a Markdown table row to have the target number of columns.

    Prepends empty cells (``| |``) to the left side of the row.
    """
    current_cols = _count_pipe_columns(line)
    if current_cols >= target_cols:
        return line

    missing = target_cols - current_cols
    stripped = line.strip()
    # Build empty cells to prepend: " |" for each missing column
    prefix = " |" * missing
    # Insert prefix right after the leading |
    if stripped.startswith("|"):
        return "|" + prefix[1:] + stripped[1:]
    return "|" + prefix[1:] + stripped


def _find_table_blocks(lines: List[str]) -> List[Tuple[int, int]]:
    """Find the start and end indices of consecutive table line blocks.

    Returns list of (start_idx, end_idx) tuples (inclusive).
    A table block is a consecutive sequence of lines starting with ``|``.
    Lines inside code blocks (``` fenced) are excluded.
    """
    blocks: List[Tuple[int, int]] = []
    in_code_block = False

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Track fenced code blocks
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # Start of a table block
        if line.startswith("|"):
            start = i
            while i < len(lines) and lines[i].strip().startswith("|"):
                # Check for code block boundary within table
                if lines[i].strip().startswith("```"):
                    break
                i += 1
            # Only count as a table if it has at least 2 lines (header + separator)
            if i - start >= 2:
                blocks.append((start, i - 1))
        else:
            i += 1

    return blocks


def fix_table_alignment(markdown: str) -> str:
    """Fix Markdown tables with inconsistent column counts.

    Detects tables where rows have different numbers of columns,
    and pads shorter rows with empty cells to match the maximum.

    Args:
        markdown: Input Markdown text containing potentially misaligned tables.

    Returns:
        Markdown text with all table rows aligned to the same column count.
    """
    if not markdown or "|" not in markdown:
        return markdown

    lines = markdown.split("\n")
    blocks = _find_table_blocks(lines)

    if not blocks:
        return markdown

    # Process blocks in reverse order so offsets don't shift
    for start, end in sorted(blocks, key=lambda b: b[0], reverse=True):
        block_lines = lines[start : end + 1]

        # Find max column count from data rows (exclude separator)
        max_cols = 0
        for line in block_lines:
            if not _is_separator_row(line):
                max_cols = max(max_cols, _count_pipe_columns(line))

        if max_cols <= 0:
            continue

        # Check if any row needs padding
        needs_fix = False
        for line in block_lines:
            if _count_pipe_columns(line) != max_cols:
                needs_fix = True
                break

        if not needs_fix:
            continue

        # Fix each row
        fixed_lines: List[str] = []
        for line in block_lines:
            if _is_separator_row(line):
                fixed_lines.append(_make_separator(max_cols))
            else:
                fixed_lines.append(_pad_row(line, max_cols))

        lines[start : end + 1] = fixed_lines

    return "\n".join(lines)


def _has_merged_cells(source: Path) -> List[bool]:
    """Check each table in a DOCX file for merged cells.

    Parses the ``<w:tbl>`` XML directly for accurate merge detection.

    Returns a list of booleans, one per table, indicating whether that
    table contains any merged cells (horizontal or vertical).

    Args:
        source: Path to the .docx file.

    Returns:
        List of booleans per table. Empty list if file cannot be read.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not available, skipping merge detection")
        return []

    try:
        doc = Document(str(source))
    except Exception as e:
        logger.debug("Cannot read DOCX for merge detection: %s", e)
        return []

    results: List[bool] = []
    for table in doc.tables:
        has_merge = False
        tbl = table._tbl
        for tr in tbl.findall(f"{{{_W_NS}}}tr"):
            for tc in tr.findall(f"{{{_W_NS}}}tc"):
                _, vmerge_status = _get_cell_grid_info(tc)
                if vmerge_status in ("restart", "continue"):
                    has_merge = True
                    break
                # Also check gridSpan > 1 via _get_cell_grid_info
                tcPr = tc.find(f"{{{_W_NS}}}tcPr")
                if tcPr is not None:
                    gs = tcPr.find(f"{{{_W_NS}}}gridSpan")
                    if gs is not None:
                        val = gs.get(f"{{{_W_NS}}}val")
                        if val and int(val) > 1:
                            has_merge = True
                            break
            if has_merge:
                break
        results.append(has_merge)

    return results


def _get_cell_grid_info(tc: ET.Element) -> Tuple[int, str]:
    """Get colspan and vMerge status from a table cell XML element.

    Args:
        tc: The ``<w:tc>`` XML element.

    Returns:
        Tuple of (colspan, vmerge_status):
        - colspan: int (1 if no gridSpan)
        - vmerge_status: "restart" | "continue" | "none"
    """
    colspan = 1
    vmerge_status = "none"

    # tcPr is a direct child of tc
    tcPr = tc.find(f"{{{_W_NS}}}tcPr")
    if tcPr is not None:
        # gridSpan inside tcPr
        gs = tcPr.find(f"{{{_W_NS}}}gridSpan")
        if gs is not None:
            val = gs.get(f"{{{_W_NS}}}val")
            if val:
                colspan = int(val)

        # vMerge inside tcPr
        vm = tcPr.find(f"{{{_W_NS}}}vMerge")
        if vm is not None:
            val = vm.get(f"{{{_W_NS}}}val")
            if val == "restart":
                vmerge_status = "restart"
            else:
                # No val attribute or val=None means continuation
                vmerge_status = "continue"

    return colspan, vmerge_status


def _build_grid_map(table) -> List[List[dict]]:
    """Build a grid map from a DOCX table using raw XML.

    Parses the ``<w:tbl>`` XML directly instead of using python-docx's
    high-level ``row.cells`` API, which returns incorrect ``_tc`` objects
    for vertically merged cells (reuses Python object IDs).

    Each grid cell is a dict with keys:
    - ``text``: cell text
    - ``colspan``: horizontal span (gridSpan)
    - ``rowspan``: vertical span (computed from vMerge chain)
    - ``skip``: True if consumed by a previous merge (omit from HTML)

    Args:
        table: A ``docx.table.Table`` object.

    Returns:
        2D list of grid cell dicts, one per logical grid position.
    """
    tbl = table._tbl
    tr_elements = tbl.findall(f"{{{_W_NS}}}tr")
    num_rows = len(tr_elements)
    if num_rows == 0:
        return []

    # First pass: determine grid width
    grid_width = 0
    if tr_elements:
        for tc in tr_elements[0].findall(f"{{{_W_NS}}}tc"):
            cs, _ = _get_cell_grid_info(tc)
            grid_width += cs
    if grid_width == 0:
        return []

    # Build occupied grid and result grid
    occupied: List[List[bool]] = [[False] * grid_width for _ in range(num_rows)]
    grid: List[List[dict]] = [[None] * grid_width for _ in range(num_rows)]

    for row_idx, tr in enumerate(tr_elements):
        tc_elements = tr.findall(f"{{{_W_NS}}}tc")
        col_idx = 0

        for tc in tc_elements:
            colspan, vmerge_status = _get_cell_grid_info(tc)
            # Extract text from <w:t> elements
            t_elements = tc.findall(f".//{{{_W_NS}}}t")
            text = "".join(t.text or "" for t in t_elements).strip()

            # For continuation cells, place at current col_idx directly
            # (they occupy the same grid position as the merge anchor above)
            if vmerge_status == "continue":
                cell_info = {
                    "text": text,
                    "colspan": colspan,
                    "rowspan": 1,
                    "skip": True,
                }
                for dc in range(colspan):
                    if col_idx + dc < grid_width:
                        occupied[row_idx][col_idx + dc] = True
                        grid[row_idx][col_idx + dc] = cell_info
                col_idx += colspan
                continue

            # For non-continue cells, skip past occupied positions first
            while col_idx < grid_width and occupied[row_idx][col_idx]:
                col_idx += 1
            if col_idx >= grid_width:
                break

            cell_info = {
                "text": text,
                "colspan": colspan,
                "rowspan": 1,
                "skip": False,
            }

            if vmerge_status == "restart":
                # Count continuation rows: look at the same column in future rows
                rowspan = 1
                for future_row_idx in range(row_idx + 1, num_rows):
                    future_tr = tr_elements[future_row_idx]
                    future_tcs = future_tr.findall(f"{{{_W_NS}}}tc")

                    # Scan future row's cells to find the one at col_idx.
                    # Process continue cells at natural position first.
                    future_col = 0
                    found_continue = False
                    for ftc in future_tcs:
                        fcs, fvm = _get_cell_grid_info(ftc)
                        if fvm == "continue":
                            if future_col == col_idx:
                                found_continue = True
                            future_col += fcs
                            continue
                        # Skip occupied for non-continue
                        while future_col < grid_width and occupied[future_row_idx][future_col]:
                            future_col += 1
                        if future_col >= grid_width:
                            break
                        future_col += fcs

                    if found_continue:
                        rowspan += 1
                        for dc in range(colspan):
                            c = col_idx + dc
                            if future_row_idx < num_rows and c < grid_width:
                                occupied[future_row_idx][c] = True
                    else:
                        break

                cell_info["rowspan"] = rowspan

            # Mark occupied positions — anchor gets full info,
            # continuation rows get skip=True (consumed by anchor)
            for dr in range(cell_info["rowspan"]):
                for dc in range(colspan):
                    r = row_idx + dr
                    c = col_idx + dc
                    if r < num_rows and c < grid_width:
                        occupied[r][c] = True
                        if dr == 0:
                            grid[r][c] = cell_info
                        else:
                            grid[r][c] = {"skip": True}

            col_idx += colspan

    return grid


def _docx_table_to_html(source: Path, table_index: int) -> Optional[str]:
    """Convert a specific table from a DOCX file to HTML.

    Properly handles both horizontal merges (colspan/gridSpan) and
    vertical merges (rowspan/vMerge) by building a complete grid map.

    Args:
        source: Path to the .docx file.
        table_index: Zero-based index of the table to convert.

    Returns:
        HTML table string, or None if conversion fails.
    """
    try:
        from docx import Document
    except ImportError:
        return None

    try:
        doc = Document(str(source))
    except Exception:
        return None

    if table_index >= len(doc.tables):
        return None

    table = doc.tables[table_index]
    num_rows = len(table.rows)
    if num_rows == 0:
        return None

    grid = _build_grid_map(table)
    if not grid:
        return None

    grid_width = len(grid[0]) if grid else 0

    rows_html: List[str] = []
    for row_idx in range(num_rows):
        cells_html: List[str] = []
        for col_idx in range(grid_width):
            cell_info = grid[row_idx][col_idx]
            if cell_info is None or cell_info.get("skip"):
                continue

            # Check if this is the "anchor" of the cell (top-left corner)
            # A cell is the anchor if it's the first occurrence in the row
            # (skip cells were already handled by the merge anchor above)
            attrs = ""
            if cell_info["colspan"] > 1:
                attrs += f' colspan="{cell_info["colspan"]}"'
            if cell_info["rowspan"] > 1:
                attrs += f' rowspan="{cell_info["rowspan"]}"'

            cells_html.append(f'<td{attrs}>{cell_info["text"]}</td>')

        if cells_html:
            rows_html.append("  <tr>" + "".join(cells_html) + "</tr>")

    if not rows_html:
        return None

    return "<table>\n" + "\n".join(rows_html) + "\n</table>"


def _find_nth_table_start(markdown: str, n: int) -> Tuple[int, int]:
    """Find the start and end line indices of the nth Markdown table.

    Args:
        markdown: The full markdown text.
        n: Zero-based table index.

    Returns:
        Tuple of (start_line, end_line) indices, or (-1, -1) if not found.
    """
    lines = markdown.split("\n")
    in_code_block = False
    table_count = 0
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        if line.startswith("|"):
            start = i
            while i < len(lines) and lines[i].strip().startswith("|"):
                if lines[i].strip().startswith("```"):
                    break
                i += 1
            end = i - 1
            if end > start:
                if table_count == n:
                    return (start, end)
                table_count += 1
        else:
            i += 1

    return (-1, -1)


def fix_merged_tables_with_html(source: Path, markdown: str) -> str:
    """Replace Markdown tables containing merged cells with HTML tables.

    For DOCX files, reads the original document structure using python-docx
    to detect merged cells (rowspan/colspan). When merges are found, the
    corresponding Markdown table is replaced with an HTML <table> that
    preserves the full cell structure.

    For non-DOCX files, returns the markdown unchanged.

    Args:
        source: Path to the original source document.
        markdown: Markdown text potentially containing tables.

    Returns:
        Markdown text with merged-cell tables replaced by HTML tables.
    """
    if not markdown:
        return markdown

    # Only process DOCX files
    if source.suffix.lower() != ".docx":
        return markdown

    # Check which tables have merged cells
    merge_flags = _has_merged_cells(source)
    if not merge_flags or not any(merge_flags):
        return markdown

    # Replace each merged table with HTML
    result = markdown
    for table_idx in range(len(merge_flags)):
        if not merge_flags[table_idx]:
            continue

        html_table = _docx_table_to_html(source, table_idx)
        if html_table is None:
            logger.debug("Could not generate HTML for table %d", table_idx)
            continue

        # Find the nth table in the current result markdown
        start, end = _find_nth_table_start(result, table_idx)
        if start == -1:
            logger.debug("Could not find table %d in markdown", table_idx)
            continue

        lines = result.split("\n")
        # Replace the table block with HTML, wrapped with blank lines
        replacement = ["", html_table, ""]
        lines[start : end + 1] = replacement
        result = "\n".join(lines)

    return result
